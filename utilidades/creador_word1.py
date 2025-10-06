from docx import Document

from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.run import WD_BREAK
from docx.shared import Pt
from docx.shared import Inches,Mm
import os
from django.conf import settings
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from io import BytesIO

class CreadorWord():
    
    def __init__(self,ruta_documento=None):        
        self.documento = Document()
        self.buffer = BytesIO()        
    
    def crearContenido(self):
        self.contenido = self.documento.add_paragraph("")
    
    def agregarTitulo(self,titulo,tipo=0):        
        if tipo == 1:
            self.runTitulo = self.contenido.add_run(f"{titulo}\n")
            self.runTitulo.bold = True
            self.runTitulo.font.name = "Segoe UI"
            self.runTitulo.underline = True
            self.runTitulo.font.size = Pt(14)
        else:
            self.runTitulo = self.contenido.add_run(f"{titulo}")
            self.runTitulo.bold = True
            self.runTitulo.font.name = "Segoe UI"
            self.runTitulo.underline = True
            self.runTitulo.font.size = Pt(14)
            
    def guardarDocumento(self,nombreDocumento):
        nombre_archivo =f"{nombreDocumento}.docx"                     
        ruta_documento = os.path.join(settings.MEDIA_ROOT,nombre_archivo)
        self.documento.save(ruta_documento)
    
    def agregarTexto(self,texto):
        self.runTexto = self.contenido.add_run(f"{texto}\n")
        self.runTexto.font.name = "Segoe UI"
        self.runTexto.font.size = Pt(10)
        
    def agregarTextoNegrita(self,texto):
        self.runTexto = self.contenido.add_run(f"{texto}")
        self.runTexto.font.bold = True
        self.runTexto.font.name = "Segoe UI"
        self.runTexto.font.size = Pt(10)       
        
    def crearTabla(self):
        self.tabla = self.documento.add_table(rows=1,cols=5)
        cabezeras = self.tabla.rows[0].cells
        cabezeras[0].text = "Nombre"
        cabezeras[1].text = "Tipo Documento"
        cabezeras[2].text = "N° Documento"
        cabezeras[3].text = "Pais"
        cabezeras[4].text = "Observación"
        
    def ingresarDataTabla(self,dataFilas):        
        nueva_fila = self.tabla.add_row().cells  
        if dataFilas["Nombre"] == "NaN" or dataFilas["Nombre"] == "nan":
            nueva_fila[0].text = ""
        else:
            nueva_fila[0].text = dataFilas["Nombre"]
            
        if dataFilas["Tipo Documento"] == "NaN" or dataFilas["Tipo Documento"] == "nan":
            nueva_fila[1].text = ""
        else:
            nueva_fila[1].text = dataFilas["Tipo Documento"]
            
        if dataFilas["N° Documento"] == "NaN" or dataFilas["N° Documento"] == "nan":
            nueva_fila[2].text = ""
        else:
            nueva_fila[2].text = dataFilas["N° Documento"]
            
        if dataFilas["Pais"] == "NaN" or dataFilas["Pais"] == "nan":
            nueva_fila[3].text = ""
        else:
            nueva_fila[3].text = dataFilas["Pais"]
            
        if dataFilas["Observación"] == "NaN" or dataFilas["Observación"] == "nan":
            nueva_fila[4].text = ""
        else:
            nueva_fila[4].text = dataFilas["Observación"]
        
    def estiloTabla(self):
        self.tabla.style = "Light Shading"
        self.tabla.autofit = True
        
    def agregarEspaciado(self):
        self.contenido.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    #SEcreto bancario Legal        
    def crear_cabezera(self):
        # Acceder a la primera sección
        section = self.documento.sections[0]
        section.header_distance = Mm(20)
        cabezera = section.header

        # Añadir imagen en el encabezado
        paragraph = cabezera.paragraphs[0] if cabezera.paragraphs else cabezera.add_paragraph()
        run = paragraph.add_run()
        ruta_imagen_cabezera = os.path.join(settings.MEDIA_ROOT,"incasur.png")
        run.add_picture(ruta_imagen_cabezera, width=Inches(1))
        cabezera.add_paragraph() 
        # Ajusta el tamaño a tu necesidad
    
    def crear_piepagina(self):    
        # Acceder a la primera sección
        section = self.documento.sections[0]        
        piepagina = section.footer

        # Añadir imagen en el encabezado
        paragraph = piepagina.paragraphs[0] if piepagina.paragraphs else piepagina.add_paragraph()
        run = paragraph.add_run("Oficina Principal: Avenida Vidaurrázaga 112-A, Parque Industrial  - Arequipa                                                           Teléfono: 054-232436")                  
        run.font.name = 'Calibri'        
        run.font.size = Pt(8)        
        
    
    def titulo_secreto_bancario(self,mes,año,correlativo):
        mi_texto = f"Carta N° {correlativo}/{mes}-{año}-SECBAN"
        texto = self.documento.add_paragraph()
        run = texto.add_run(mi_texto)
        run.font.name = "Arial Narrow"
        run.font.size = Pt(11)  
        texto.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run.add_break(WD_BREAK.LINE)
        run.add_break(WD_BREAK.LINE)
            
    def añadir_parte1(self,texto):
        p = self.documento.add_paragraph()
        p.paragraph_format.line_spacing = Pt(14)
         
        texto_completo = f"SEÑORES:\n{texto}"       
        run1 = p.add_run(texto_completo)
        run1.font.name = "Arial Narrow"
        run1.font.size = Pt(11)  
        run1.bold = True                
        
    def parrafo_contenido_princial(self):
        self.parrafo_principal = self.documento.add_paragraph()
        self.parrafo_principal.paragraph_format.line_spacing = Pt(23)        
        self.parrafo_principal.space_after = Pt(0)
    
    def agregar_contenidos(self,texto1,texto2,tipo):     
        if tipo == 1:
            run3 = self.parrafo_principal.add_run(texto1)
            run3.font.name = "Arial Narrow"
            run3.font.size = Pt(11)  
            run3.bold = True
            run4 = self.parrafo_principal.add_run(f"{texto2}\n")
            run4.font.name = "Arial Narrow"
            run4.font.size = Pt(11)
        else:
            run3 = self.parrafo_principal.add_run(texto1)
            run3.font.name = "Arial Narrow"
            run3.font.size = Pt(11)  
            run3.bold = True
            run4 = self.parrafo_principal.add_run(texto2)
            run4.font.name = "Arial Narrow"
            run4.font.size = Pt(11)

    def escrito_final1(self):
        parrafo = self.documento.add_paragraph()
        texto_completo = "CAJA RURAL DE AHORRO Y CRÉDITO INCASUR S.A."
        run1 = parrafo.add_run(texto_completo)
        run1.font.name = "Arial Narrow"
        run1.font.size = Pt(11)  
        run1.bold = True

        texto_completo = ", inscrita en la Partida Registral Nº 11179010 del Registro de Personas Jurídicas de la Oficina Registral de Arequipa, con domicilio en Av. Vidaurrázaga 112-A, Parque Industrial, distrito, provincia y departamento de Arequipa, a usted decimos:\t"
        run2 = parrafo.add_run(texto_completo)
        run2.font.name = "Arial Narrow"
        run2.font.size = Pt(11) 
    
    def escrito_final2(self):
        parrafo = self.documento.add_paragraph()
        texto_completo = "Las siguientes personas naturales o jurídicas,"
        run1 = parrafo.add_run(texto_completo)
        run1.font.name = "Arial Narrow"
        run1.font.size = Pt(11)  

        texto_completo = "no mantienen ninguna relación comercial o contractual con nuestra entidad:\n"
        run2 = parrafo.add_run(texto_completo)
        run2.font.name = "Arial Narrow"
        run2.font.size = Pt(11) 
        run2.bold = True
        run2.underline = True
        
    def tabla_secreto_bancario(self):   
       # --- Configuración de Estilos ---
        COLOR_AZUL_OSCURO = '3A567D' # Fondo del encabezado (HEX)
        COLOR_BLANCO = RGBColor(0xFF, 0xFF, 0xFF) # Fuente del encabezado
        
        # --- 1. Definición de Anchos Fijos ---
        # La suma es 6.7 pulgadas. Asegúrate de que este ancho total quepa en tus márgenes.
        ANCHO_COLUMNAS = [
            Inches(0.4),
            Inches(3.2),
            Inches(1.5),
            Inches(1.6)
            ]
        self.ANCHO_COLUMNAS = ANCHO_COLUMNAS # Almacenar para usar en añadir_fila_sb
        
        # 2. Creación y Configuración Básica de la Tabla
        self.tabla = self.documento.add_table(rows=1, cols=4)
        self.tabla.style = 'Table Grid'
        self.tabla.autofit = False # Clave: Desactivar ajuste automático

        # --- 3. Forzar el Ancho de la Tabla al 100% del Contenido ---
        # Esto ayuda a la compatibilidad en Word, estirando la tabla al ancho de página.
        tbl = self.tabla._tbl
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:type'), 'pct') # Tipo: porcentaje
        tblW.set(qn('w:w'), '5000')
        tbl.tblPr.append(tblW)
        
        # 4. Aplicar Anchos Fijos a las Columnas (a nivel de columna, por si acaso)
        for i, width in enumerate(ANCHO_COLUMNAS):
            self.tabla.columns[i].width = width

        # --- 5. Estilizado del Encabezado y FORZAR Ancho Fijo a Nivel de Celda ---
        cabezeras = self.tabla.rows[0].cells
        titulos = ["N°", "Nombres y Apellidos", "Tipo Documento", "N° Documento"]

        for i, cell in enumerate(cabezeras):
            
            # 🌟 PASO CLAVE: Forzar Ancho Fijo en la CELDA (para Word) 🌟
            ancho_twips = ANCHO_COLUMNAS[i].twips
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:type'), 'dxa') # Tipo: dxa (medida fija)
            tcW.set(qn('w:w'), str(ancho_twips))
            
            # Reemplazar el elemento tcW existente
            for existing_tcW in tcPr.xpath('./w:tcW'):
                tcPr.remove(existing_tcW)
            tcPr.append(tcW)
            
            # A. Aplicar Sombreado de Fondo (Azul Oscuro)
            tcShd = OxmlElement('w:shd')
            tcShd.set(qn('w:fill'), COLOR_AZUL_OSCURO) 
            tcPr.append(tcShd)

            # B. Asignar Texto y Formato
            paragraph = cell.paragraphs[0]
            paragraph.text = titulos[i]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Formato del texto (Negrita y Blanco)
            run = paragraph.runs[0]
            run.bold = True
            run.font.color.rgb = COLOR_BLANCO
            run.font.size = Pt(10)
        
    def añadir_fila_sb(self,dataFilas):        
        nueva_fila = self.tabla.add_row().cells
    
        # Usar la lista de anchos guardada en el objeto 'self'
        ANCHO_COLUMNAS = self.ANCHO_COLUMNAS 
        
        # Iterar sobre las celdas para aplicar texto, alineación y ancho fijo
        for i, cell in enumerate(nueva_fila):
            
            # 🌟 PASO CLAVE: Forzar Ancho Fijo en la CELDA (Repetido aquí) 🌟
            ancho_twips = ANCHO_COLUMNAS[i].twips
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:type'), 'dxa') # Tipo: dxa (medida fija)
            tcW.set(qn('w:w'), str(ancho_twips))
            
            # Reemplazar el elemento tcW existente
            for existing_tcW in tcPr.xpath('./w:tcW'):
                tcPr.remove(existing_tcW)
            tcPr.append(tcW)
            
            # --- Lógica de Asignación de Contenido ---
            
            # Simplificación de la lógica de NaN/nan
            data_key = ["Contador", "Nombre", "Tipo Documento", "N° Documento"][i]
            
            contenido = dataFilas.get(data_key, "")
            if isinstance(contenido, str) and contenido.lower() in ("nan"):
                cell.text = ""
            else:
                cell.text = str(contenido)

            # Establecer Alineación
            # Columna N° (0) y N° Documento (3) centradas, el resto a la izquierda (por defecto)
            if i == 0 or i == 3:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def agregar_texto_normal(self):
        parrafo = self.documento.add_paragraph()        
        texto_completo = "\nSin otro particular,"
        run1 = parrafo.add_run(texto_completo)
        run1.font.name = "Arial Narrow"
        run1.font.size = Pt(11)
        
    def agregar_texto_derecha(self,dia,mes,año):
        meses = {'01':'enero','02':'febrero','03':'marzo','04':'abril','05':'mayo','06':'junio','07':'julio','08':'agosto','09':'setiembre','10':'octubre','11':'noviembre','12':'diciembre'}
        mes_texto = meses[str(mes)]
        mi_texto = f"Arequipa, {dia} de {mes_texto} de {año}"
        texto = self.documento.add_paragraph()
        run = texto.add_run(mi_texto)
        run.font.name = "Arial Narrow"
        run.font.size = Pt(11)  
        texto.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        
    def salto_pagina(self):
        paragraph = self.documento.add_paragraph() # Crea un nuevo párrafo
        run = paragraph.add_run()             # Crea un nuevo objeto Run dentro del párrafo
        run.add_break(WD_BREAK.PAGE)
                       
        
# Crear un nuevo documento
# document = Document()

# # Añadir un título al documento
# contenido = document.add_paragraph("")
# runTitulo1 = contenido.add_run("Datos de Identificacion y Ubicacion :\n")
# runTitulo1.bold = True
# runTitulo1.font.name = "Segoe UI"
# runTitulo1.underline = True
# runTitulo1.font.size = Pt(14)

# titulos_data = ["Número de envio:","Fecha de Envio de Paquete:","Tipo de solicitud:","Número de Expediente:","Entidad Solicitante:","Nombre de la Autoridad:","Número de oficio de la autoridad:","Dirección de la Autoridad:","Delito / Materia:","Información requerida:","Información adicional:","N° Expediente/Carpeta Fiscal/Caso:","Prioridad Alta:","Plazo de respuesta:","Periodo de consulta:"]
# for td in titulos_data:
#     run = contenido.add_run(f"{td}\n")    
#     run.font.name = "Segoe UI"
    
# runTitulo2 = contenido.add_run("Personas incluidas en la solicitud :")
# runTitulo2.bold = True
# runTitulo2.font.name = "Segoe UI"
# runTitulo2.underline = True
# runTitulo2.font.size = Pt(14)

# contenido.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

# tabla = document.add_table(rows=1,cols=5)
# cabezeras = tabla.rows[0].cells
# cabezeras[0].text = "Nombre"
# cabezeras[1].text = "Tipo Documento"
# cabezeras[2].text = "N° Documento"
# cabezeras[3].text = "Pais"
# cabezeras[4].text = "Observación"

# #Aplicas estilo a la cabezera
# for celda in cabezeras:
#     celda.paragraphs[0].runs[0].bold = True
#     celda.paragraphs[0].runs[0].font.size = Pt(13)
# #Estilos a la Tabla
# tabla.style = "Light Shading"
# tabla.autofit = True

# datos_a_añadir = [
#     {"Nombre": "Juan Pérez", "Tipo Documento": "DNI", "N° Documento": "12345678", "País": "Perú", "Observación": "Cliente nuevo"},
#     {"Nombre": "Maria García", "Tipo Documento": "Pasaporte", "N° Documento": "P9876543", "País": "España", "Observación": "Pendiente de verificación"},
#     {"Nombre": "Carlos Rojas", "Tipo Documento": "CE", "N° Documento": "E11223344", "País": "Colombia", "Observación": ""},
#     {"Nombre": "Ana Gómez", "Tipo Documento": "DNI", "N° Documento": "99887766", "País": "México", "Observación": "VIP"},
# ]

# #Rellena las filas
# for filas in datos_a_añadir:    
#     nueva_fila = tabla.add_row().cells
    
#     nueva_fila[0].text = filas["Nombre"]
#     nueva_fila[1].text = filas["Tipo Documento"]
#     nueva_fila[2].text = filas["N° Documento"]
#     nueva_fila[3].text = filas["País"]
#     nueva_fila[4].text = filas["Observación"]
        
# document.save("documento_con_tabla_ejemplo.docx")

# print("Documento 'documento_con_tabla_ejemplo.docx' creado exitosamente con una tabla.")

